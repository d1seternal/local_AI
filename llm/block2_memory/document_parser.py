"""
Модуль для обработки PDF- и DOCX-документов.

"""

import pathlib
from typing import List, Dict, Any
from dataclasses import dataclass

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

from docling.document_converter import DocumentConverter, PdfFormatOption
from docling.datamodel.base_models import InputFormat
from docling.datamodel.pipeline_options import PdfPipelineOptions, TableFormerMode

try:
    import pandas as pd
    from tabulate import tabulate
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False


@dataclass
class Chunk:
    text: str
    metadata: Dict[str, Any]


@dataclass
class DocumentProcessingResult:
    text: str                    
    chunks: List[Chunk]         
    tables: List[Dict]            
    metadata: Dict[str, Any]   


class DocumentProcessor:
    def __init__(
        self,
        use_docling: bool = True,
        ocr_enabled: bool = False,
        table_mode: str = "fast",      
        chunk_size: int = 200,        
        chunk_overlap: int = 50
    ):
        self.use_docling = use_docling
        self.ocr_enabled = ocr_enabled
        self.table_mode = table_mode
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        if self.use_docling:
            self._init_docling()
    
    def _init_docling(self):
        pipeline_options = PdfPipelineOptions()
        pipeline_options.do_ocr = self.ocr_enabled
        pipeline_options.do_table_structure = True
        pipeline_options.generate_picture_images = False
        pipeline_options.generate_page_images = False
        pipeline_options.images_scale = 1.0
        
        if self.table_mode == "accurate":
            pipeline_options.table_structure_options.mode = TableFormerMode.ACCURATE
        else:
            pipeline_options.table_structure_options.mode = TableFormerMode.FAST
        
        pipeline_options.do_code_enrichment = False
        
        self.docling_converter = DocumentConverter(
            format_options={
                InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
            }
        )
    
    def process_document(self, file_path: pathlib.Path) -> DocumentProcessingResult:
        ext = file_path.suffix.lower()
    
        if ext == '.pdf' and self.use_docling:
            return self._process_pdf(file_path)
     
        elif ext == '.docx' and DOCX_SUPPORT:
            return self._process_docx(file_path)
        else:
            return self._process_fallback(file_path)
    
    def _process_pdf(self, file_path: pathlib.Path) -> DocumentProcessingResult:
        try:
            result = self.docling_converter.convert(str(file_path))
            document = result.document
            data = document.export_to_dict()
            
        
            all_text_parts = []
            tables_info = []
         
            texts = data.get('texts', [])
            for text_item in texts:
                text = text_item.get('text', '').strip()
                if text:
                    all_text_parts.append(text)
          
            tables = data.get('tables', [])
            for table_idx, table_data in enumerate(tables):
                try:
                    table = document.tables[table_idx]
                    table_md = self._table_to_md(table.model_dump())
                 
                    context = self._find_table_context(table_data, data)
                    
                    table_text = []
                    if context:
                        table_text.append(f"[Таблица: {context}]")
                    table_text.append(table_md)
                    
                    all_text_parts.append("\n".join(table_text))
                    
                    tables_info.append({
                        'id': table_idx,
                        'text': table_md,
                        'context': context
                    })
                    
                except Exception as e:
                    print(f"Ошибка таблицы {table_idx}: {e}")

            full_text = "\n\n".join(all_text_parts)
            chunks = self._create_chunks(full_text)
            metadata = {
                "filename": file_path.name,
                "file_size": file_path.stat().st_size,
                "tables_count": len(tables_info),
                "parser": "docling",
                "ocr_used": self.ocr_enabled,
                "table_mode": self.table_mode
            }
            
            return DocumentProcessingResult(
                text=full_text,
                chunks=chunks,
                tables=tables_info,
                metadata=metadata
            )
            
        except Exception as e:
            print(f"Ошибка PDF: {e}")
            return self._process_fallback(file_path)
    
    def _process_docx(self, file_path: pathlib.Path) -> DocumentProcessingResult:
        try:
            doc = DocxDocument(file_path)
            
            all_text_parts = []
            tables_info = []

            for para in doc.paragraphs:
                if para.text.strip():
                    all_text_parts.append(para.text.strip())
 
            for table_idx, table in enumerate(doc.tables):
                try:
                    table_lines = []
                    for row in table.rows:
                        row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_cells:
                            table_lines.append(" | ".join(row_cells))
                    
                    if table_lines:
                        table_text = "\n".join(table_lines)
                        all_text_parts.append(f"[ТАБЛИЦА {table_idx+1}]\n{table_text}")
                        
                        tables_info.append({
                            'id': table_idx,
                            'text': table_text,
                            'rows': len(table.rows),
                            'cols': len(table.columns) if hasattr(table, 'columns') else 0
                        })
                except Exception as e:
                    print(f"Ошибка таблицы {table_idx}: {e}")

            full_text = "\n\n".join(all_text_parts)
            chunks = self._create_chunks(full_text)    
            metadata = {
                "filename": file_path.name,
                "file_size": file_path.stat().st_size,
                "tables_count": len(tables_info),
                "paragraphs_count": len(doc.paragraphs),
                "parser": "python-docx"
            }
            
            return DocumentProcessingResult(
                text=full_text,
                chunks=chunks,
                tables=tables_info,
                metadata=metadata
            )
            
        except Exception as e:
            print(f"Ошибка DOCX: {e}")
            return self._process_fallback(file_path)
    
    def _process_txt(self, file_path: pathlib.Path) -> DocumentProcessingResult:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='cp1251') as f:
                text = f.read()
        
        chunks = self._create_chunks(text)
        
        return DocumentProcessingResult(
            text=text,
            chunks=chunks,
            tables=[],
            metadata={
                "filename": file_path.name,
                "file_size": file_path.stat().st_size,
                "parser": "txt"
            }
        )
    
    def _process_fallback(self, file_path: pathlib.Path) -> DocumentProcessingResult:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return DocumentProcessingResult(
                text=text,
                chunks=self._create_chunks(text),
                tables=[],
                metadata={
                    "filename": file_path.name,
                    "file_size": file_path.stat().st_size,
                    "parser": "fallback"
                }
            )
        except:
            return DocumentProcessingResult(
                text="",
                chunks=[],
                tables=[],
                metadata={
                    "filename": file_path.name,
                    "error": "Failed to read file"
                }
            )
    
    def _find_table_context(self, table_data: Dict, data: Dict) -> str:
        for text_item in data.get('texts', []):
            text = text_item.get('text', '').strip()
            if any(word in text.lower() for word in ['таблица', 'table', 'рис', 'fig']):
                return text[:100]
        return ""
    
    def _table_to_md(self, table: Dict) -> str:
        if not PANDAS_AVAILABLE:
            return "[Таблица]"
        
        try:
            grid = table.get('data', {}).get('grid', [])
            if not grid:
                return "[Пустая таблица]"
            
            table_data = []
            for row in grid:
                table_row = [cell.get('text', '') for cell in row]
                table_data.append(table_row)
            
            if len(table_data) > 1:
                headers = table_data[0]
                data = table_data[1:]
                return tabulate(data, headers=headers, tablefmt="github")
            else:
                return tabulate(table_data, tablefmt="github")
                
        except Exception:
            return "[Ошибка таблицы]"
    
    def _create_chunks(self, text: str) -> List[Chunk]:
        words = text.split()
        if len(words) <= self.chunk_size:
            return [Chunk(text=text, metadata={"chunk_index": 0})]
        
        chunks = []
        step = self.chunk_size - self.chunk_overlap
        
        for i in range(0, len(words), step):
            chunk_words = words[i:i + self.chunk_size]
            if chunk_words:
                chunk_text = " ".join(chunk_words)
                chunks.append(Chunk(
                    text=chunk_text,
                    metadata={"chunk_index": i // step}
                ))
        
        return chunks

def process_document(file_path: pathlib.Path, **kwargs) -> DocumentProcessingResult:
    processor = DocumentProcessor(**kwargs)
    return processor.process_document(file_path)