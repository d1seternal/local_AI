"""
Модуль для обработки PDF- и DOCX-документов.

"""
import json
import pathlib
from pathlib import Path
import traceback
from typing import List, Dict, Any
from dataclasses import dataclass
from langchain_text_splitters import RecursiveCharacterTextSplitter
import re


import warnings
warnings.filterwarnings("ignore", category=DeprecationWarning, module="PyPDF2")

try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    print("поддержка PDF отключена")

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

from docling.document_converter import DocumentConverter, PdfFormatOption
from docling.datamodel.base_models import InputFormat
from docling.datamodel.pipeline_options import PdfPipelineOptions, TableFormerMode

try:
    import pandas as pd
    from tabulate import tabulate
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

from shared.config import CHUNK_SIZE, CHUNK_OVERLAP

@dataclass
class Chunk:
    text: str
    metadata: Dict[str, Any]


@dataclass
class DocumentProcessingResult:
    text: str                    
    chunks: List[Chunk]         
    tables: List[Dict]            
    metadata: Dict[str, Any]   


class DocumentProcessor:
    def __init__(
        self,
        use_docling: bool = True,
        ocr_enabled: bool = True,
        table_mode: str = "fast",      
        chunk_size = CHUNK_SIZE,        
        chunk_overlap = CHUNK_OVERLAP
        # max_num_pages: int = 0
    ):
        self.use_docling = use_docling
        self.ocr_enabled = ocr_enabled
        self.table_mode = table_mode
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        # self.max_num_pages = max_num_pages
        
        if self.use_docling:
            self._init_docling()
    
    MAX_DOCLING_PAGES = 15

    def _init_docling(self):
        self._base_pipeline_options = PdfPipelineOptions()
        self._base_pipeline_options.do_ocr = self.ocr_enabled
        self._base_pipeline_options.do_table_structure = True
        self._base_pipeline_options.generate_picture_images = False
        self._base_pipeline_options.generate_page_images = False
        self._base_pipeline_options.images_scale = 1.0
        # self._base_pipeline_options.max_num_pages = self.MAX_DOCLING_PAGES
        
        if self.table_mode == "accurate":
            self._base_pipeline_options.table_structure_options.mode = TableFormerMode.ACCURATE
        else:
            self._base_pipeline_options.table_structure_options.mode = TableFormerMode.FAST
        
        self._base_pipeline_options.do_code_enrichment = False
        
        self.docling_converter = DocumentConverter(
            format_options={
                InputFormat.PDF: PdfFormatOption(pipeline_options=self._base_pipeline_options)
            }
        )

    def process_document(self, file_path: pathlib.Path) -> DocumentProcessingResult:
        ext = file_path.suffix.lower()
    
        if ext == '.pdf' and self.use_docling:
            return self._process_pdf(file_path)
     
        elif ext == '.docx' and DOCX_SUPPORT:
            return self._process_docx(file_path)
        else:
            return self._process_fallback(file_path)

    def _get_pdf_page_count(self, file_path: pathlib.Path) -> int:
        if not PDF_SUPPORT:
            return 999
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                return len(reader.pages)
        except Exception:
            return 999
    
    def _process_pdf(self, file_path: pathlib.Path) -> DocumentProcessingResult:
        
        page_count = self._get_pdf_page_count(file_path)
        
        print(f"\n PDF: {file_path.name}")
        print(f"   Страниц: {page_count}")
        
        if page_count <= self.MAX_DOCLING_PAGES: 
            if not self.use_docling:
                print(f"  Docling не доступен, используем PyPDF2")
                return self._process_pdf_with_pypdf2_enhanced(file_path)
            
            try:
                print(f"  Запуск Docling...")
                result = self._process_pdf_with_docling(file_path, ocr_enabled=True)
                print(f"  Docling завершен успешно")
                return result
            except MemoryError as e:
                print(f"  Ошибка памяти в Docling: {e}")
                print(f"  Переключаемся на PyPDF2...")
                return self._process_pdf_with_pypdf2_enhanced(file_path)
            except Exception as e:
                print(f"   Ошибка Docling: {e}")
                print(f"   Переключаемся на PyPDF2...")
                return self._process_pdf_with_pypdf2_enhanced(file_path)
        else:
            return self._process_pdf_with_pypdf2_enhanced(file_path)
    
    def _process_pdf_with_docling(self, file_path: pathlib.Path, ocr_enabled) -> DocumentProcessingResult:
        try:
            result = self.docling_converter.convert(str(file_path))
            document = result.document
            data = document.export_to_dict()
            all_text_parts = []
            tables_info = []
            
            texts = data.get('texts', [])
            for text_item in texts:
                text = text_item.get('text', '').strip()
                if text:
                    all_text_parts.append(text)

            tables = data.get('tables', [])
            for table_idx, table_data in enumerate(tables):
                try:
                    table = document.tables[table_idx]
                    table_md = self._table_to_md(table.model_dump())
                    
                    context = self._find_table_context(table_data, data)
                    
                    table_text = []
                    if context:
                        table_text.append(f"[Таблица: {context}]")
                    table_text.append(table_md)
                    
                    all_text_parts.append("\n".join(table_text))
                    
                    tables_info.append({
                        'id': table_idx,
                        'text': table_md,
                        'context': context
                    })
                except Exception as e:
                    print(f"   Ошибка таблицы {table_idx}: {e}")
            
            full_text = "\n\n".join(all_text_parts)
            chunks = self._create_chunks(full_text)
            
            return DocumentProcessingResult(
                text=full_text,
                chunks=chunks,
                tables=tables_info,
                metadata={
                    "filename": file_path.name,
                    "parser": "docling",
                    "ocr_used": ocr_enabled,
                    "tables_count": len(tables_info),
                    "pages_processed": min(self._get_pdf_page_count(file_path), self.MAX_DOCLING_PAGES)
                }
            )
            
        except Exception as e:
            raise Exception(f"Docling processing failed: {e}")
    
    def _process_pdf_with_pypdf2_enhanced(self, file_path: pathlib.Path) -> DocumentProcessingResult:
        if not PDF_SUPPORT:
            return self._process_fallback(file_path)
        
        try:
            all_text = []
            tables_info = []
            
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                total_pages = len(reader.pages)
                
                print(f"  Обработка через PyPDF2")
                
                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text:
                        page_text, page_tables = self._enhance_text_with_tables(text)
                        all_text.append(page_text)
                        tables_info.extend(page_tables)
                    
                    if (page_num + 1) % 10 == 0:
                        print(f"  Обработано страниц: {page_num + 1}/{total_pages}")
            
            full_text = "\n\n".join(all_text)
            chunks = self._create_chunks(full_text)
            
            return DocumentProcessingResult(
                text=full_text,
                chunks=chunks,
                tables=tables_info,
                metadata={
                    "filename": file_path.name,
                    "file_size": file_path.stat().st_size,
                    "pages": total_pages,
                    "tables_count": len(tables_info),
                    "parser": "pypdf2_enhanced"
                }
            )
            
        except Exception as e:
            print(f"  PyPDF2 ошибка: {e}")
            return self._process_fallback(file_path)
    
    def _enhance_text_with_tables(self, text: str) -> tuple:
        lines = text.split('\n')
        tables = []
        current_table = []
        in_table = False
        table_id = 0
        
        for line in lines:
            # Признаки таблицы: множественные пробелы (3+), разделители, выравнивание
            is_table_line = False
            
            # Множественные пробелы (вероятный признак таблицы)
            if re.search(r'\s{3,}', line):
                is_table_line = True
            
            # Разделители таблиц
            if '|' in line or re.search(r'[+-]{3,}', line):
                is_table_line = True
            
            # Много чисел в строке (характерно для таблиц)
            numbers = re.findall(r'\d+', line)
            if len(numbers) >= 3:
                is_table_line = True
            
            if is_table_line:
                if not in_table:
                    in_table = True
                    current_table = []
                    table_id += 1
                
                # Форматируем строку таблицы
                if '|' in line:
                    # Уже есть разделители
                    formatted_line = line
                else:
                    # Разделяем по множественным пробелам
                    cells = re.split(r'\s{2,}', line.strip())
                    if len(cells) >= 2:
                        formatted_line = "| " + " | ".join(cells) + " |"
                    else:
                        formatted_line = line
                
                current_table.append(formatted_line)
            else:
                if in_table and len(current_table) >= 2:
                    # Добавляем разделитель заголовка
                    first_row_cells = current_table[0].split('|')[1:-1]
                    if first_row_cells:
                        header_separator = "|" + "|".join(["---" for _ in first_row_cells]) + "|"
                        current_table.insert(1, header_separator)
                    
                    tables.append({
                        'id': table_id,
                        'text': '\n'.join(current_table),
                        'rows': len(current_table) - 2  # минус заголовок и разделитель
                    })
                
                current_table = []
                in_table = False
        
        # Обработка последней таблицы
        if in_table and len(current_table) >= 2:
            first_row_cells = current_table[0].split('|')[1:-1]
            if first_row_cells:
                header_separator = "|" + "|".join(["---" for _ in first_row_cells]) + "|"
                current_table.insert(1, header_separator)
            
            tables.append({
                'id': table_id,
                'text': '\n'.join(current_table),
                'rows': len(current_table) - 2
            })
        
        # Вставляем маркеры таблиц в текст
        formatted_text = text
        for table in tables:
            # Добавляем маркер таблицы
            table_marker = f"\n[ТАБЛИЦА {table['id']}]\n{table['text']}\n"
            formatted_text = formatted_text.replace(table['text'], table_marker)
        
        return formatted_text, tables
    
    
    def _process_docx(self, file_path: pathlib.Path) -> DocumentProcessingResult:
        try:
            doc = DocxDocument(file_path) 
            all_content = []
            tables_info = []

            for para in doc.paragraphs:
                if para and para.text.strip():
                    clean_text = para.text.strip()
                    clean_text = re.sub(r'[\x00-\x1f\x7f]', '', clean_text)
                    clean_text = re.sub(r'\s+', ' ', clean_text)
                    if clean_text:
                        all_content.append(clean_text)
            
            for table_idx, table in enumerate(doc.tables):
                try:
                    table_lines = []
                    for row in table.rows:
                        row_cells = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                cell_text = re.sub(r'[\x00-\x1f\x7f]', '', cell_text)
                                cell_text = re.sub(r'\s+', ' ', cell_text)
                                row_cells.append(cell_text)
                        
                        if row_cells:
                            table_lines.append(" | ".join(row_cells))
                    
                    if table_lines:
                        table_text = "\n".join(table_lines)
                        all_content.append(f"[ТАБЛИЦА {table_idx+1}]")
                        all_content.append(table_text)
                        
                        tables_info.append({
                            'id': table_idx,
                            'text': table_text,
                            'rows': len(table.rows),
                            'cols': len(table.columns) if hasattr(table, 'columns') else 0
                        })
                except Exception as e:
                    print(f"Ошибка таблицы {table_idx}: {e}")
            
            full_text = "\n\n".join(all_content)
            if full_text:
                full_text = re.sub(r'<[^>]+>', '', full_text)
                full_text = re.sub(r'[\x00-\x1f\x7f]', '', full_text)
                full_text = re.sub(r'\s+', ' ', full_text)
                full_text = re.sub(r'[^\w\s\u0400-\u04FF.,!?;:()\-–—]', ' ', full_text)
                full_text = re.sub(r'\s+', ' ', full_text)
                full_text = full_text.strip()
            
            if not full_text or len(full_text) < 50:
                return DocumentProcessingResult(
                    text="",
                    chunks=[],
                    tables=[],
                    metadata={
                        "filename": file_path.name,
                        "error": "Текст слишком короткий или отсутствует"
                    }
                )
            
            print(f"Извлечено текста: {len(full_text)} символов")
            
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=CHUNK_SIZE,          
                chunk_overlap=CHUNK_OVERLAP,    
                length_function=len,
                separators=["\n\n", "\n", ". ", "! ", "? ", ", ", " ", ""] 
            )
            
            chunks = text_splitter.split_text(full_text)
            
            print(f"Создано {len(chunks)} чанков")

            chunk_objects = []
            for idx, chunk_text in enumerate(chunks):
                clean_chunk = ' '.join(chunk_text.split())
                if len(clean_chunk) > 20:
                    chunk_objects.append(
                        Chunk(
                            text=clean_chunk,
                            metadata={"chunk_index": idx}
                        )
                    )
            
            print(f"DOCX обработан: {len(chunks)} чанков, {len(tables_info)} таблиц\n")
            
            return DocumentProcessingResult(
                text=full_text,
                chunks=chunk_objects,
                tables=tables_info,
                metadata={
                    "filename": file_path.name,
                    "file_size": file_path.stat().st_size,
                    "tables_count": len(tables_info),
                    "parser": "python-docx"
                }
            )
        
        except Exception as e:
            traceback.print_exc()
            return DocumentProcessingResult(
                text="",
                chunks=[],
                tables=[],
                metadata={
                    "filename": file_path.name,
                    "error": str(e)
                }
            )  
    
    def _process_txt(self, file_path: pathlib.Path) -> DocumentProcessingResult:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='cp1251') as f:
                text = f.read()
        
        chunks = self._create_chunks(text)
        
        return DocumentProcessingResult(
            text=text,
            chunks=chunks,
            tables=[],
            metadata={
                "filename": file_path.name,
                "file_size": file_path.stat().st_size,
                "parser": "txt"
            }
        )
    
    def _process_fallback(self, file_path: pathlib.Path) -> DocumentProcessingResult:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return DocumentProcessingResult(
                text=text,
                chunks=self._create_chunks(text),
                tables=[],
                metadata={
                    "filename": file_path.name,
                    "file_size": file_path.stat().st_size,
                    "parser": "fallback"
                }
            )
        except:
            return DocumentProcessingResult(
                text="",
                chunks=[],
                tables=[],
                metadata={
                    "filename": file_path.name,
                    "error": "Failed to read file"
                }
            )
    
    def _find_table_context(self, table_data: Dict, data: Dict) -> str:
        for text_item in data.get('texts', []):
            text = text_item.get('text', '').strip()
            if any(word in text.lower() for word in ['таблица', 'table', 'рис', 'fig']):
                return text[:100]
        return ""
    
    def _table_to_md(self, table: Dict) -> str:
        if not PANDAS_AVAILABLE:
            return "[Таблица]"
        
        try:
            grid = table.get('data', {}).get('grid', [])
            if not grid:
                return "[Пустая таблица]"
            
            table_data = []
            for row in grid:
                table_row = [cell.get('text', '') for cell in row]
                table_data.append(table_row)
            
            if len(table_data) > 1:
                headers = table_data[0]
                data = table_data[1:]
                return tabulate(data, headers=headers, tablefmt="github")
            else:
                return tabulate(table_data, tablefmt="github")
                
        except Exception:
            return "[Ошибка таблицы]"
    
    def _create_chunks(self, text: str) -> List[Chunk]:
        words = text.split()
        if len(words) <= self.chunk_size:
            return [Chunk(text=text, metadata={"chunk_index": 0})]
        
        chunks = []
        step = self.chunk_size - self.chunk_overlap
        
        for i in range(0, len(words), step):
            chunk_words = words[i:i + self.chunk_size]
            if chunk_words:
                chunk_text = " ".join(chunk_words)
                chunks.append(Chunk(
                    text=chunk_text,
                    metadata={"chunk_index": i // step}
                ))
        
        return chunks

def process_document(file_path: pathlib.Path, **kwargs) -> DocumentProcessingResult:
    processor = DocumentProcessor(**kwargs)
    return processor.process_document(file_path)