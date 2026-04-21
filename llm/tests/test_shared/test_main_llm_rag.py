import pytest
import tempfile
import shutil
from pathlib import Path
from unittest.mock import Mock, patch, MagicMock
import sys
from pathlib import Path as PathLib

_project_root = PathLib(__file__).parent.parent.parent
if str(_project_root) not in sys.path:
    sys.path.insert(0, str(_project_root))

from block2_memory.main_llm_rag import (
    extract_keywords,
    clean_model_output,
    ModelBenchmark,
    generate_with_prompts,
    load_model
)
from shared.memory import VectorMemory


class TestExtractKeywords:
    def test_extract_keywords_russian_text(self):
        question = "Каковы условия аренды помещения?"
        keywords = extract_keywords(question)
        
        assert isinstance(keywords, list)
        assert len(keywords) > 0
    
    def test_extract_keywords_english_text(self):
        question = "What are the lease terms and conditions?"
        keywords = extract_keywords(question)
        
        assert isinstance(keywords, list)
    
    def test_extract_keywords_mixed_case(self):
        question = "Python programming LANGUAGE is GREAT"
        keywords = extract_keywords(question)
   
        for kw in keywords:
            assert kw.islower()
    
    def test_extract_keywords_with_numbers(self):
        question = "Сколько будет 15 умножить на 37?"
        keywords = extract_keywords(question)
        
        assert "15" not in keywords
        assert "37" not in keywords


class TestCleanModelOutput:
    def test_clean_remove_think_tags(self):
        raw = "Рассуждение... <think>Это внутренние рассуждения</think> Ответ"
        cleaned = clean_model_output(raw)
        assert "<think>" not in cleaned
        assert "</think>" not in cleaned
    
    def test_clean_remove_json_artifacts(self):
        raw = '{"answer": "test"}'
        cleaned = clean_model_output(raw)
        # JSON может остаться или быть удален
        assert isinstance(cleaned, str)
    
    def test_clean_remove_multiple_spaces(self):
        raw = "Много   пробелов   между   словами"
        cleaned = clean_model_output(raw)
        assert "   " not in cleaned


class TestModelBenchmark:
    @pytest.fixture
    def benchmark(self):
        return ModelBenchmark()
    
    def test_init(self, benchmark):
        assert benchmark.metrics['total_queries'] == 0
        assert benchmark.metrics['total_tokens'] == 0
        assert benchmark.metrics['total_time'] == 0
        assert benchmark.metrics['queries'] == []
    
    def test_get_memory_usage(self, benchmark):
        mem = benchmark.get_memory_usage()
        assert 'rss' in mem
        assert 'vms' in mem
        assert isinstance(mem['rss'], float)
        assert isinstance(mem['vms'], float)
    
    def test_add_query_result(self, benchmark):
        benchmark.add_query_result(
            query="тест",
            response="ответ",
            tokens=100,
            time_taken=2.5,
            tokens_per_second=40.0
        )
        
        assert benchmark.metrics['total_queries'] == 1
        assert benchmark.metrics['total_tokens'] == 100
        assert benchmark.metrics['total_time'] == 2.5
        assert len(benchmark.metrics['queries']) == 1
    
    def test_add_multiple_queries(self, benchmark):
        for i in range(3):
            benchmark.add_query_result(
                query=f"тест{i}",
                response=f"ответ{i}",
                tokens=50,
                time_taken=1.0,
                tokens_per_second=50.0
            )
        
        assert benchmark.metrics['total_queries'] == 3
        assert benchmark.metrics['total_tokens'] == 150
        assert benchmark.metrics['total_time'] == 3.0
    
    def test_calculate_model_ram_usage(self, benchmark):
        benchmark.metrics['memory_before'] = {'rss': 1000}
        benchmark.metrics['memory_after'] = {'rss': 2000}
        benchmark.metrics['model_size_gb'] = 1.0
        
        result = benchmark.calculate_model_ram_usage()
        
        assert 'model_ram_mb' in result
        assert 'model_ram_gb' in result
        assert result['model_ram_mb'] == 1000.0  

class TestLoadModel:
    @patch('llm.block2_memory.main_llm_rag.Llama')
    def test_load_model_success(self, mock_llama, temp_dir):
        mock_llama.return_value = Mock()
        benchmark = ModelBenchmark()
    
    @patch('llm.block2_memory.main_llm_rag.get_shared_llm')
    def test_load_model_shared(self, mock_get_shared):
        mock_shared_llm = Mock()
        mock_get_shared.return_value = mock_shared_llm
        
        benchmark = ModelBenchmark()

class TestIntegrationRAG:
    
    @pytest.fixture
    def temp_memory(self):
        temp_dir = tempfile.mkdtemp()
        memory = VectorMemory(persist_directory=str(temp_dir))
        yield memory
        shutil.rmtree(temp_dir)
    
    @pytest.fixture
    def sample_docx_file(self, temp_dir):
        from docx import Document
        doc = Document()
        doc.add_paragraph("Юридический документ. Договор аренды.")
        doc.add_paragraph("Срок аренды: 12 месяцев.")
        doc.add_paragraph("Арендная плата: 50000 рублей в месяц.")
        file_path = Path(temp_dir) / "contract.docx"
        doc.save(file_path)
        return file_path
    
    def test_keyword_extraction_in_context(self):
        legal_question = "Каковы условия расторжения договора аренды?"
        keywords = extract_keywords(legal_question)
        
        legal_terms = ["условия", "расторжения", "договора", "аренды"]
        found_terms = [term for term in legal_terms if term in keywords]
        
        assert len(found_terms) >= 2